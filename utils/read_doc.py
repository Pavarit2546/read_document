from docx import Document
import json
import io

def read_docx_content(docx_file_stream: bytes) -> str:
    """
    ฟังก์ชันสำหรับอ่านและดึงข้อความทั้งหมดจากไฟล์ DOCX
    
    :param docx_file_stream: Byte Stream ของไฟล์ DOCX
    :return: JSON string ที่มีข้อความทั้งหมดในเอกสาร
    """
    try:
        # ใช้ io.BytesIO เพื่อให้ Document สามารถอ่านจาก Stream ได้
        document = Document(io.BytesIO(docx_file_stream))
        
        full_text = []
        # อ่านข้อความจากย่อหน้าทั้งหมด
        for paragraph in document.paragraphs:
            full_text.append(paragraph.text)
            
        # อ่านข้อความจากตารางทั้งหมด
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        text_content = "\n".join(full_text)
        
        output = {
            "status": "success",
            "extracted_text": text_content
        }
        # ส่งออกเป็น JSON string 
        return json.dumps(output, ensure_ascii=False)
    
    except Exception as e:
        return json.dumps({"status": "error", "message": f"Error reading DOCX: {str(e)}"}, ensure_ascii=False)